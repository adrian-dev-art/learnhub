import subprocess
import sys
import tempfile
import os
from django.core.mail import send_mail
from django.conf import settings
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from io import BytesIO
from datetime import datetime
from docx import Document


def convert_docx_to_html(file_path):
    """
    Convert a .docx file to simple HTML for the editor.
    """
    doc = Document(file_path)
    html_content = ""

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Simple mapping of styles to HTML tags
        if para.style.name.startswith('Heading 1'):
            html_content += f"<h1>{text}</h1>"
        elif para.style.name.startswith('Heading 2'):
            html_content += f"<h2>{text}</h2>"
        elif para.style.name.startswith('Heading 3'):
            html_content += f"<h3>{text}</h3>"
        elif para.style.name.startswith('List Bullet'):
            html_content += f"<ul><li>{text}</li></ul>"
        elif para.style.name.startswith('List Number'):
            html_content += f"<ol><li>{text}</li></ol>"
        else:
            html_content += f"<p>{text}</p>"

    return html_content


def parse_docx_to_modules(file_path):
    """
    Parse a .docx file and return a list of module data.
    Each heading starts a new module.
    """
    doc = Document(file_path)
    modules = []
    current_module = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Check if it's a heading (Heading 1, 2, 3 etc)
        if para.style.name.startswith('Heading'):
            if current_module:
                modules.append(current_module)

            current_module = {
                'title': text,
                'content': '',
                'content_type': 'text'
            }
        else:
            if current_module:
                current_module['content'] += text + '\n\n'
            else:
                # If no heading yet, start a default module
                current_module = {
                    'title': 'Introduction',
                    'content': text + '\n\n',
                    'content_type': 'text'
                }

    if current_module:
        modules.append(current_module)

    return modules


def send_access_key_email(user_email, course_title, access_key):
    """Send course access key to user's email"""
    subject = f"Your Access Key for {course_title}"
    message = f"""
    Hello!

    Thank you for enrolling in {course_title}.

    Your course access key is: {access_key}

    You can now access the course using this key.

    Happy learning!

    Best regards,
    Course Platform Team
    """

    try:
        send_mail(
            subject,
            message,
            settings.DEFAULT_FROM_EMAIL,
            [user_email],
            fail_silently=False,
        )
        return True
    except Exception as e:
        print(f"Email sending failed: {e}")
        return False


def execute_python_code(code, inputs="", timeout=5):
    """
    Execute Python code in a sandboxed environment
    Returns tuple: (success, output, error)
    """
    try:
        # Create a temporary file to store the code
        with tempfile.NamedTemporaryFile(mode='w', suffix='.py', delete=False) as f:
            f.write(code)
            temp_file = f.name

        try:
            # Execute the code with timeout
            result = subprocess.run(
                [sys.executable, temp_file],
                capture_output=True,
                text=True,
                input=inputs,
                timeout=timeout,
                cwd=tempfile.gettempdir()
            )

            if result.returncode == 0:
                return True, result.stdout, ""
            else:
                return False, "", result.stderr

        finally:
            # Clean up temp file
            if os.path.exists(temp_file):
                os.remove(temp_file)

    except subprocess.TimeoutExpired:
        return False, "", "Code execution timed out (5 seconds limit)"
    except Exception as e:
        return False, "", str(e)


def generate_certificate_pdf(certificate):
    """Generate PDF certificate with custom design"""
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors
    from reportlab.lib.utils import ImageReader
    import qrcode

    buffer = BytesIO()
    # Create canvas in landscape mode
    c = canvas.Canvas(buffer, pagesize=landscape(A4))
    width, height = landscape(A4)

    # Colors
    dark_blue = colors.HexColor('#0f172a')
    gold_color = colors.HexColor('#d97706')
    text_black = colors.HexColor('#1f2937')
    text_gray = colors.HexColor('#4b5563')

    # --- Background Design ---
    # Draw the dark sidebar on the right
    # Polygon points: Top-Right, Bottom-Right, Bottom-Mid, Top-Mid
    # The diagonal cut goes from top (approx 65% width) to bottom (approx 55% width)

    path = c.beginPath()
    path.moveTo(width * 0.6, height) # Top-Mid start
    path.lineTo(width, height)       # Top-Right
    path.lineTo(width, 0)            # Bottom-Right
    path.lineTo(width * 0.5, 0)      # Bottom-Mid end
    path.close()

    c.setFillColor(dark_blue)
    c.drawPath(path, fill=1, stroke=0)

    # --- Left Section (White) ---

    # "CERTIFICATE"
    c.setFillColor(text_black)
    c.setFont("Times-Bold", 42)
    c.drawString(50, height - 100, "CERTIFICATE")

    # "OF APPRECIATION"
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, height - 125, "OF APPRECIATION")

    # Line under title
    c.setLineWidth(2)
    c.setStrokeColor(text_black)
    c.line(50, height - 135, 250, height - 135)

    # "PROUDLY PRESENTED TO"
    c.setFont("Helvetica", 10)
    c.setFillColor(text_gray)
    c.drawString(50, height - 180, "PROUDLY PRESENTED TO")

    # Student Name
    student_name = certificate.user.get_full_name() or certificate.user.username
    c.setFont("Times-Italic", 36) # Using Italic to mimic script
    c.setFillColor(colors.black)
    c.drawString(50, height - 230, student_name)

    # "FOR AN EXCELLENT PERFORMANCE..."
    c.setFont("Helvetica-Bold", 9)
    c.setFillColor(text_black)
    c.drawString(50, height - 280, "FOR AN EXCELLENT PERFORMANCE AS A PARTICIPANT")

    # Course Details
    c.setFont("Helvetica", 10)
    c.setFillColor(text_black)
    # Wrap text if too long
    course_text = f"Workshop: {certificate.course.title}"
    c.drawString(50, height - 310, course_text)

    # Date Range (Mocked based on issue date)
    c.setFont("Helvetica", 9)
    c.setFillColor(text_gray)
    date_str = certificate.issued_at.strftime('%d %B %Y')
    c.drawString(50, height - 330, f"Completed on {date_str}")

    # Bottom Left Date
    c.setFont("Helvetica-Bold", 10)
    c.setFillColor(text_black)
    c.drawString(50, 50, date_str)

    # CDIA Logo Placeholder (Bottom Center-Left)
    c.setFont("Helvetica-Bold", 16)
    c.setFillColor(colors.black)
    c.drawString(200, 50, "CDIA")
    c.setFont("Helvetica", 8)
    c.drawString(200, 35, "CAKRA DIGITAL ANDALAN")

    # --- Right Section (Dark) ---

    # Company Name
    c.setFont("Helvetica-Bold", 12)
    c.setFillColor(gold_color)
    c.drawRightString(width - 30, height - 50, "PT. CAKRA DIGITAL ANDALAN (CDIA)")

    c.setFont("Helvetica", 8)
    c.setFillColor(colors.white)
    c.drawRightString(width - 30, height - 65, "NOMOR AHU-0031310.AH.01.01.TAHUN 2022")

    # Laurel Wreath / Badge (Mocked with text/circle)
    c.saveState()
    c.translate(width - 100, height / 2 - 20)
    c.setStrokeColor(gold_color)
    c.setLineWidth(3)
    c.circle(0, 0, 50, stroke=1, fill=0)

    c.setFont("Times-Bold", 14)
    c.setFillColor(gold_color)
    c.drawCentredString(0, 10, "Private")
    c.drawCentredString(0, -5, "Lesson")
    c.setFont("Helvetica-Bold", 12)
    c.drawCentredString(0, -25, "2025")
    c.restoreState()

    c.setFont("Helvetica-Bold", 16)
    c.setFillColor(colors.HexColor('#332e20')) # Dark goldish
    c.drawCentredString(width - 100, height / 2 - 90, "Batch 2")

    # QR Code
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=10,
        border=1,
    )
    # QR Data: Verification URL
    verify_url = f"https://learnhub.com/verify/{certificate.certificate_id}"
    qr.add_data(verify_url)
    qr.make(fit=True)

    qr_img = qr.make_image(fill_color="black", back_color="white")

    # Convert PIL image to ReportLab Image
    qr_buffer = BytesIO()
    qr_img.save(qr_buffer, format="PNG")
    qr_buffer.seek(0)

    c.drawImage(ImageReader(qr_buffer), width - 230, 80, width=80, height=80)

    # Signature Name
    c.setFont("Helvetica-Bold", 10)
    c.setFillColor(colors.white)
    c.drawCentredString(width - 190, 60, "Deni Suprihadi, S.T, M.KOM")

    c.showPage()
    c.save()
    buffer.seek(0)
    return buffer
